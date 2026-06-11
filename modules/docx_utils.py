import re
import win32com.client as win32
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from modules import file_utils
from modules.constants import HandleResult
import os

# pip install python-docx pywin32

POINTS_PER_CM = 72 / 2.54
DEFAULT_WATERMARK_LEFT = 4
DEFAULT_WATERMARK_TOP = 6
DEFAULT_WATERMARK_WIDTH = 4
DEFAULT_WATERMARK_HEIGHT = None

WD_RELATIVE_HORIZONTAL_POSITION_PAGE = 1
WD_RELATIVE_VERTICAL_POSITION_PAGE = 1
WD_WRAP_BEHIND = 5
WD_BORDER_BOTTOM = -3
WD_LINE_STYLE_NONE = 0


def cm_to_points(value):
    return value * POINTS_PER_CM


def docx_to_pdf(docx_path, pdf_path):
    """
    将指定的 DOCX 文件转换为 PDF 文件。

    :param docx_path: 输入的 DOCX 文件的路径
    :param pdf_path: 输出的 PDF 文件的路径
    """
    result = HandleResult.Created
    if os.path.exists(pdf_path):
        if file_utils.compare_file_modify_time(docx_path, pdf_path) == 1:
            # 源文件比目标文件新，则删除目标文件
            os.remove(pdf_path)
            result = HandleResult.Overrided
        else:
            # 源文件比目标文件旧，则跳过
            result = HandleResult.Skiped
            return result

    # 创建 Word 应用程序对象
    word = win32.DispatchEx("Word.Application")
    try:
        # 打开 DOCX 文件
        doc = word.Documents.Open(docx_path)
        # 将文档保存为 PDF 格式
        doc.SaveAs(pdf_path, FileFormat=17)
        # 关闭文档
        doc.Close()
    finally:
        if word is not None:
            # 退出 Word 应用程序
            word.Quit()
    return result


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def split_table_cells(line):
    text = line.strip()
    if "\t" in text:
        cells = re.split(r"\t+", text)
    elif "|" in text:
        cells = text.strip("|").split("|")
    elif re.search(r" {2,}", text):
        cells = re.split(r" {2,}", text)
    else:
        return None

    cells = [cell.strip() for cell in cells]
    if len(cells) < 2 or any(not cell for cell in cells):
        return None
    return cells


def set_cell_black_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_black_borders(cell)


def add_txt_lines(doc, lines):
    index = 0
    while index < len(lines):
        cells = split_table_cells(lines[index])
        if cells is None:
            doc.add_paragraph(lines[index])
            index += 1
            continue

        rows = [cells]
        col_count = len(cells)
        next_index = index + 1
        while next_index < len(lines):
            next_cells = split_table_cells(lines[next_index])
            if next_cells is None or len(next_cells) != col_count:
                break
            rows.append(next_cells)
            next_index += 1

        if len(rows) >= 2:
            add_table(doc, rows)
        else:
            doc.add_paragraph(lines[index])
        index = next_index


def txt_to_docx(txt_path, docx_path=None):
    txt_path = os.path.abspath(txt_path)
    docx_path = docx_path or os.path.splitext(txt_path)[0] + ".docx"
    result = HandleResult.Overrided if os.path.exists(docx_path) else HandleResult.Created

    with open(txt_path, "r", encoding="utf-8") as file:
        lines = [line.strip() for line in file if line.strip()]

    doc = Document(docx_path) if os.path.exists(docx_path) else Document()
    clear_body(doc)
    add_txt_lines(doc, lines)
    doc.save(docx_path)
    doc = None
    return result


def add_image_watermark(
    docx_path,
    image_path,
    left=DEFAULT_WATERMARK_LEFT,
    top=DEFAULT_WATERMARK_TOP,
    width=DEFAULT_WATERMARK_WIDTH,
    height=None,
):
    """
    使用 pywin32 库为指定的 DOCX 文件添加指定图片水印。

    :param docx_path: 要处理的 DOCX 文件的路径
    :param image_path: 水印图片的路径
    :param left: 水印图片的水平位置，单位为 cm
    :param top: 水印图片的垂直位置，单位为 cm
    :param width: 水印图片宽度，单位为 cm
    :param height: 水印图片高度，单位为 cm。为 None 时按宽度等比自适应
    """
    left_points = cm_to_points(left)
    top_points = cm_to_points(top)
    width_points = cm_to_points(width)
    height_points = None if height is None else cm_to_points(height)

    # 创建 Word 应用程序对象
    word = win32.DispatchEx("Word.Application")
    # 不显示 Word 应用程序窗口
    word.Visible = False
    try:
        # 打开 DOCX 文件
        doc = word.Documents.Open(docx_path)
        # 获取活动文档的节
        sections = doc.Sections
        for section in sections:
            # 获取页眉
            header = section.Headers(1)
            # 激活页眉进行编辑
            header.Range.Select()
            # 添加图片到页眉
            shape = header.Shapes.AddPicture(
                FileName=image_path,
                LinkToFile=False,
                SaveWithDocument=True,
                Left=left_points,
                Top=top_points,
            )
            shape.LockAspectRatio = True
            shape.Width = width_points
            if height_points is not None:
                shape.Height = height_points
            shape.RelativeHorizontalPosition = WD_RELATIVE_HORIZONTAL_POSITION_PAGE
            shape.RelativeVerticalPosition = WD_RELATIVE_VERTICAL_POSITION_PAGE
            # 设置图片的环绕方式为衬于文字下方
            shape.WrapFormat.Type = WD_WRAP_BEHIND
            # 设置图片的透明度 (实际测试不起作用)
            shape.Fill.Transparency = 0.5
            # 去除图片的填充（灰色背景）
            shape.Fill.Visible = 0
            # 去除页眉分隔线
            header.Range.Borders(WD_BORDER_BOTTOM).LineStyle = WD_LINE_STYLE_NONE            

        # 保存文档
        doc.Save()
        doc.Close()
    finally:
        if word is not None:
            # 退出 Word 应用程序
            word.Quit()
    return HandleResult.Processed
